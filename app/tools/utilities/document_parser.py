# document_parser.py
import abc
import os
import logging
from typing import Optional
import io
import boto3
from urllib.parse import urlparse
import fitz  # PyMuPDF
from pathlib import Path
from bs4 import BeautifulSoup
import markdown
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)
logging.basicConfig(level=logging.INFO)

class DocumentParser(abc.ABC):
    @abc.abstractmethod
    def parse(self, filepath: str) -> str:
        """Extract text from filepath and return it."""
        pass

class DOCXParser(DocumentParser):
    def parse(self, filepath: str) -> str:
        logger.info(f"Parsing DOCX: {filepath}")
        s3_helper = S3Helper()
        try:
            if s3_helper.is_s3_path(filepath):
                bio = io.BytesIO(s3_helper.fetch_s3_bytes(filepath))
                doc = DocxDocument(bio)
            else:
                doc = DocxDocument(filepath)
            return "\n\n".join(p.text for p in doc.paragraphs)
        except Exception:
            logger.exception(f"Error parsing DOCX: {filepath}")
            raise

class PDFParser(DocumentParser):
    def parse(self, filepath: str) -> str:
        logger.info(f"Parsing PDF: {filepath}")
        s3_helper = S3Helper()
        try:
            if s3_helper.is_s3_path(filepath):
                bio = io.BytesIO(s3_helper.fetch_s3_bytes(filepath))
                doc = fitz.open(stream=bio, filetype="pdf")
            else:
                doc = fitz.open(filepath)
            return "\n".join(page.get_text() for page in doc)
        except Exception as e:
            logger.exception(f"Error parsing PDF: {filepath} - Message: {str(e)}")
            raise

class TXTParser(DocumentParser):
    def parse(self, filepath: str) -> str:
        logger.info(f"Parsing TXT: {filepath}")
        s3_helper = S3Helper()
        try:
            if s3_helper.is_s3_path(filepath):
                data = s3_helper.fetch_s3_bytes(filepath).decode('utf-8')
            else:
                with open(filepath, 'r', encoding='utf-8') as f:
                    data = f.read()
            return data
        except Exception:
            logger.exception(f"Error reading TXT: {filepath}")
            raise

class MDParser(DocumentParser):
    def parse(self, filepath: str) -> str:
        logger.info(f"Parsing Markdown: {filepath}")
        s3_helper = S3Helper()
        try:
            if s3_helper.is_s3_path(filepath):
                logger.debug("Detected S3 URI, fetching bytes")
                raw = s3_helper.fetch_s3_bytes(filepath).decode("utf-8")
            else:
                raw = Path(filepath).read_text(encoding="utf-8")

            html = markdown.markdown(raw)
            soup = BeautifulSoup(html, "html.parser")
            return "".join(soup.strings)

        except FileNotFoundError:
            logger.exception(f"Markdown file not found: {filepath}")
            raise
        except Exception as e:
            logger.exception(f"Error parsing Markdown: {e}")
            raise


class DocumentParserFactory:
    _parsers = {
        '.pdf': PDFParser(),
        '.docx': DOCXParser(),
        '.txt': TXTParser(),
        '.md': MDParser(),
    }

    @staticmethod
    def get_parser(filepath: str) -> Optional[DocumentParser]:
        _, ext = os.path.splitext(filepath.lower())
        return DocumentParserFactory._parsers.get(ext)
    

class S3Helper:
    _s3_client = boto3.client('s3')
    
    def is_s3_path(self, path: str) -> bool:
        return path.lower().startswith('s3://')
    
    def fetch_s3_bytes(self, s3_path: str) -> bytes:
        """Download the S3 object and return its content as bytes."""
        parsed = urlparse(s3_path)
        bucket = parsed.netloc
        key = parsed.path.lstrip('/')
        obj = self._s3_client.get_object(Bucket=bucket, Key=key)
        return obj['Body'].read()


def parse_document(filepath: str) -> str:
    parser = DocumentParserFactory.get_parser(filepath)
    if not parser:
        raise ValueError(f"No parser available for: {filepath}")
    return parser.parse(filepath)

